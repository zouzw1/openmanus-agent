#!/usr/bin/env python3
"""
学术论文格式化工具

按照学术论文标准格式化 Word 文档。

使用方法：
    uv run --with python-docx python3 format_academic.py input.docx output.docx
"""

import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def apply_academic_format(doc):
    """应用学术论文格式标准"""

    ref_pattern = re.compile(r'\[\d+\]')

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # 论文标题（第一段）
        if i == 0 and len(text) < 50:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(18)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(18)  # 小二
                run.font.bold = True

        # 一级标题（第X章 或 X. ）
        elif (text.startswith('第') and '章' in text) or re.match(r'^\d+\.\s', text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(15)  # 小三
                run.font.bold = True

        # 二级标题（X.X ）
        elif re.match(r'^\d+\.\d+\s', text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(14)  # 四号
                run.font.bold = True

        # 三级标题（X.X.X ）
        elif re.match(r'^\d+\.\d+\.\d+\s', text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = 1.5
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(12)  # 小四
                run.font.bold = True

        # 参考文献（[1] [2] ...）
        elif text.startswith('[') and ']' in text[:5]:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(-21)
            para.paragraph_format.left_indent = Pt(21)
            para.paragraph_format.line_spacing = 1.0
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)  # 五号

        # 图片说明（图X 或 Figure X）
        elif text.startswith('图') or text.startswith('Figure'):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing = 1.0
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(10.5)  # 五号

        # 正文（长段落）
        elif len(text) > 30:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(24)  # 2字符
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)  # 小四

                # 引用编号变蓝色
                if ref_pattern.search(run.text):
                    run.font.color.rgb = RGBColor(0, 0, 255)

    # 处理表格
    for table in doc.tables:
        # 表头（第一行）
        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                        run.font.size = Pt(10.5)  # 五号
                        run.font.bold = True

        # 表格内容（其余行）
        for row in table.rows[1:]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)  # 五号

    # 设置页面（A4纸，标准边距）
    section = doc.sections[0]
    section.page_height = Pt(845)  # A4高度
    section.page_width = Pt(598)   # A4宽度
    section.top_margin = Pt(72)    # 2.54cm (1英寸)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)   # 3.17cm
    section.right_margin = Pt(90)


def main():
    if len(sys.argv) < 3:
        print("用法: uv run --with python-docx python3 format_academic.py input.docx output.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    print(f"正在加载文档: {input_path}")
    doc = Document(input_path)

    print("正在应用学术论文格式标准...")
    apply_academic_format(doc)

    print(f"正在保存文档: {output_path}")
    doc.save(output_path)

    print("✓ 格式化完成！")
    print("\n格式标准:")
    print("  - 标题: 黑体，小二，居中")
    print("  - 正文: 宋体/Times New Roman，小四，1.5倍行距")
    print("  - 首行缩进: 2字符(24pt)")
    print("  - 参考文献: 悬挂缩进，引用编号蓝色")
    print("  - 页边距: 上下2.54cm，左右3.17cm")


if __name__ == "__main__":
    main()
