#!/usr/bin/env python3
"""
Word 文档格式分析工具

分析 Word 文档的字体、字号、行距、缩进等格式信息。

使用方法：
    uv run --with python-docx python3 analyze.py input.docx
"""

import sys
from docx import Document
from collections import defaultdict


def emu_to_pt(emu):
    """EMU 转 pt"""
    return round(emu / 914400 * 72, 1) if emu else None


def analyze_document(docx_path):
    """分析文档格式"""
    doc = Document(docx_path)

    # 统计数据
    size_stats = defaultdict(int)
    font_stats = defaultdict(int)
    indent_stats = defaultdict(int)
    line_spacing_stats = defaultdict(int)

    def analyze_para(para):
        """分析单个段落"""
        for run in para.runs:
            if run.font.size:
                size_stats[emu_to_pt(run.font.size)] += 1
            if run.font.name:
                font_stats[run.font.name] += 1

        # 段落格式
        if para.paragraph_format.first_line_indent:
            indent = emu_to_pt(para.paragraph_format.first_line_indent)
            indent_stats[indent] += 1

        if para.paragraph_format.line_spacing:
            spacing = para.paragraph_format.line_spacing
            if isinstance(spacing, int):
                line_spacing_stats[f"{spacing}倍"] += 1
            else:
                line_spacing_stats[f"{emu_to_pt(spacing)}pt"] += 1

    # 分析普通段落
    for para in doc.paragraphs:
        analyze_para(para)

    # 分析表格内段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    analyze_para(para)

    # 输出结果
    print(f"\n{'='*60}")
    print(f"文档格式分析报告: {docx_path}")
    print(f"{'='*60}\n")

    print("字号分布:")
    for size, count in sorted(size_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {size}pt: {count}次")

    print("\n字体分布:")
    for font, count in sorted(font_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {font}: {count}次")

    print("\n首行缩进分布:")
    for indent, count in sorted(indent_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {indent}pt: {count}次")

    print("\n行距分布:")
    for spacing, count in sorted(line_spacing_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {spacing}: {count}次")

    print(f"\n{'='*60}\n")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: uv run --with python-docx python3 analyze.py input.docx")
        sys.exit(1)

    analyze_document(sys.argv[1])
