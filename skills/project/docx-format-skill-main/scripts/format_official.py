#!/usr/bin/env python3
"""
公文格式化工具

按照 GB/T 9704-2012《党政机关公文格式》标准格式化 Word 文档。

使用方法：
    uv run --with python-docx python3 format_official.py input.docx output.docx
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # 创建边框元素
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')  # 实线
        edge_element.set(qn('w:sz'), '4')  # 0.5磅 = 4 (1/8 pt)
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')  # 黑色
        tcBorders.append(edge_element)

    tcPr.append(tcBorders)


def apply_official_format(doc):
    """应用公文格式标准"""

    # 1. 处理标题（第一段）
    if doc.paragraphs:
        title_para = doc.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(18)
        for run in title_para.runs:
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(22)  # 二号
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色

    # 2. 处理正文和标题
    for i, para in enumerate(doc.paragraphs[1:], 1):
        text = para.text.strip()
        if not text:
            continue

        # 一级标题（一、二、三...）
        if text.startswith('一、') or text.startswith('二、') or text.startswith('三、') or \
           text.startswith('四、') or text.startswith('五、') or text.startswith('六、') or \
           text.startswith('七、') or text.startswith('八、') or text.startswith('九、') or \
           text.startswith('十、'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.first_line_indent = Pt(0)  # 标题无缩进
            for run in para.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)  # 三号
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        # 二级标题（（一）（二）...）
        elif text.startswith('（一）') or text.startswith('（二）') or text.startswith('（三）') or \
             text.startswith('（四）') or text.startswith('（五）') or text.startswith('（六）'):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.first_line_indent = Pt(0)  # 标题无缩进
            for run in para.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体_GB2312')
                run.font.size = Pt(16)  # 三号
                run.font.color.rgb = RGBColor(0, 0, 0)

        # 三级标题（1.1、1.2...）
        elif len(text) < 50 and ('.' in text[:10] or '：' in text or ':' in text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = Pt(28)
            para.paragraph_format.first_line_indent = Pt(0)  # 标题无缩进
            for run in para.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(16)  # 三号
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        # 正文（所有其他段落）
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = Pt(32)  # 2字符
            para.paragraph_format.line_spacing = Pt(28)  # 固定28磅
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                run.font.size = Pt(16)  # 三号
                run.font.color.rgb = RGBColor(0, 0, 0)

    # 3. 处理表格
    for table in doc.tables:
        # 表格居中
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 设置表格居中对齐
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        tblPr.append(jc)

        # 处理表格内容
        for row in table.rows:
            for cell in row.cells:
                # 设置边框
                set_cell_border(cell)

                # 格式化单元格内容
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.first_line_indent = Pt(0)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                        run.font.size = Pt(10.5)  # 五号
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # 4. 设置页面（A4纸，公文标准边距）
    section = doc.sections[0]
    section.page_height = Pt(845)  # A4高度
    section.page_width = Pt(598)   # A4宽度
    section.top_margin = Pt(106)   # 3.7cm
    section.bottom_margin = Pt(85) # 3.0cm
    section.left_margin = Pt(99)   # 3.5cm
    section.right_margin = Pt(85)  # 3.0cm


def main():
    if len(sys.argv) < 3:
        print("用法: uv run --with python-docx python3 format_official.py input.docx output.docx")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    print(f"正在加载文档: {input_path}")
    doc = Document(input_path)

    print("正在应用公文格式标准...")
    apply_official_format(doc)

    print(f"正在保存文档: {output_path}")
    doc.save(output_path)

    print("✓ 格式化完成！")
    print("\n格式标准:")
    print("  - 标题: 黑体，二号，居中，黑色")
    print("  - 正文: 仿宋_GB2312，三号，固定行距28磅，黑色")
    print("  - 首行缩进: 2字符(32pt)")
    print("  - 表格: 居中，实线边框(0.5磅)，黑色")
    print("  - 页边距: 上3.7cm，下3.0cm，左3.5cm，右3.0cm")


if __name__ == "__main__":
    main()
