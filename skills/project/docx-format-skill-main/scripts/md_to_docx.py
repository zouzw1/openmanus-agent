#!/usr/bin/env python3
"""Markdown 转 DOCX 工具

直接读取 Markdown 文档并按格式要求写入 DOCX，无需 pandoc。

使用方法：
    uv run --with python-docx python3 md_to_docx.py input.md output.docx
"""

import sys
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_font(run, cn='宋体', en='Times New Roman', size=12, bold=False):
    """设置中英文字体"""
    run.font.name = en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=12):
    """添加带格式的文本（支持粗体）"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, cn=cn, en=en, size=size, bold=True)
        elif part:
            run = para.add_run(part)
            set_font(run, cn=cn, en=en, size=size, bold=False)


def add_heading(doc, text, level):
    """添加标题"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        set_font(run, cn='黑体', en='Arial', size=15, bold=True)
    elif level == 2:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        set_font(run, cn='黑体', en='Arial', size=14, bold=True)
    elif level == 3:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        set_font(run, cn='黑体', en='Arial', size=12, bold=True)
    elif level == 4:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        set_font(run, cn='黑体', en='Arial', size=12, bold=True)


def add_paragraph(doc, text):
    """添加正文段落"""
    if not text.strip():
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=12)


def add_list_item(doc, text, ordered=False, level=0, restart_numbering=False, number=None):
    """添加列表项"""
    if ordered and number is not None:
        # 手动添加序号
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        if level > 0:
            para.paragraph_format.left_indent = Pt(24 * level)
        # 添加序号和文本
        full_text = f"{number}. {text}"
        add_text_with_formatting(para, full_text, cn='宋体', en='Times New Roman', size=12)
    else:
        para = doc.add_paragraph(style='List Bullet' if not ordered else 'List Number')
        para.paragraph_format.line_spacing = 1.5
        if level > 0:
            para.paragraph_format.left_indent = Pt(24 * level)

        # 重启有序列表编号
        if ordered and restart_numbering:
            from docx.oxml import OxmlElement
            pPr = para._element.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')
            ilvl = OxmlElement('w:ilvl')
            ilvl.set(qn('w:val'), str(level))
            numPr.append(ilvl)
            numId = OxmlElement('w:numId')
            numId.set(qn('w:val'), '1')
            numPr.append(numId)
            pPr.append(numPr)

        add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=12)


def add_table(doc, lines, start_idx):
    """添加表格"""
    table_lines = []
    i = start_idx

    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        # 跳过分隔行（包含 - : 等字符的行）
        if line and not re.match(r'^\|[\s\-:|\s]+\|$', line):
            table_lines.append(line)
        i += 1

    if not table_lines:
        return start_idx

    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)

    if rows:
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))

        # 设置表格边框为实线黑色 0.5 磅
        from docx.oxml import OxmlElement
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 0.5磅 = 4/8磅
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # 填充表格内容
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text_with_formatting(para, cell_text, cn='宋体', en='Times New Roman', size=10.5)

    return i - 1


def parse_markdown(md_file, doc):
    """解析 Markdown 文件"""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    prev_type = None  # 跟踪上一个段落类型
    prev_level = -1   # 跟踪上一个列表的缩进级别

    while i < len(lines):
        line = lines[i].rstrip()

        # 空行跳过
        if not line:
            prev_type = 'empty'
            i += 1
            continue

        # 分隔线跳过
        if line.strip() == '---':
            prev_type = 'separator'
            i += 1
            continue

        # 标题（从长到短匹配）
        if line.startswith('#### '):
            add_heading(doc, line[5:], 4)
            prev_type = 'heading'
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
            prev_type = 'heading'
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
            prev_type = 'heading'
        elif line.startswith('# '):
            add_heading(doc, line[2:], 1)
            prev_type = 'heading'

        # 表格
        elif line.startswith('|'):
            i = add_table(doc, lines, i)
            prev_type = 'table'

        # 任务列表（- [ ] 或 - [x]）
        elif re.match(r'^\s*-\s+\[([ xX])\]\s+', line):
            text = re.sub(r'^\s*-\s+\[([ xX])\]\s+', '', line)
            indent_level = (len(line) - len(line.lstrip())) // 2
            add_list_item(doc, text, ordered=False, level=indent_level)
            prev_type = 'list_bullet'
            prev_level = indent_level

        # 无序列表（支持缩进）
        elif re.match(r'^\s*[-*]\s+', line):
            indent_level = (len(line) - len(line.lstrip())) // 2
            text = re.sub(r'^\s*[-*]\s+', '', line)
            add_list_item(doc, text, ordered=False, level=indent_level)
            prev_type = 'list_bullet'
            prev_level = indent_level

        # 有序列表（支持缩进）
        elif re.match(r'^\s*\d+\.\s', line):
            indent_level = (len(line) - len(line.lstrip())) // 2
            # 提取序号
            match = re.match(r'^\s*(\d+)\.\s', line)
            number = int(match.group(1))
            text = re.sub(r'^\s*\d+\.\s', '', line)
            # 使用原始序号
            add_list_item(doc, text, ordered=True, level=indent_level, number=number)
            prev_type = 'list_number'
            prev_level = indent_level

        # 代码块（跳过）
        elif line.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                i += 1
            prev_type = 'code'

        # 正文
        else:
            add_paragraph(doc, line)
            prev_type = 'paragraph'

        i += 1


def main():
    if len(sys.argv) != 3:
        print("使用方法: python3 md_to_docx.py input.md output.docx")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    doc = Document()
    parse_markdown(input_file, doc)
    doc.save(output_file)

    print(f"✓ 转换完成: {input_file} -> {output_file}")


if __name__ == '__main__':
    main()
