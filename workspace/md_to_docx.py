from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

# 创建新文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
# 设置中文字体
font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 读取Markdown内容（简化版处理）
md_content = '''# Java（初学者）8 周学习路线图\n\n## 总体安排\n- 总周期：8 周\n- 每周投入：10 小时\n- 推荐节奏：2 次概念学习 + 1 次编码练习 + 1 次复盘总结\n\n## 第 1 周：环境搭建与基础语法\n### 学习重点\n- 理解 JDK、JRE、JVM 的关系\n- 完成开发环境搭建与项目运行\n- 掌握变量、数据类型、输入输出\n\n### 本周练习\n1. 编写 HelloWorld 程序并成功运行\n2. 完成变量声明、类型转换和四则运算程序\n\n## 第 2 周：流程控制与数组\n### 学习重点\n- 掌握 if/else、switch 分支语句\n- 理解 for、while、do-while 循环\n- 学习一维数组的声明、初始化和遍历\n\n### 本周练习\n1. 实现九九乘法表\n2. 编写成绩评级程序\n3. 开发猜数字游戏\n\n## 第 3 周：面向对象编程基础\n### 学习重点\n- 理解类与对象的概念\n- 掌握封装、构造方法、this 关键字\n- 学习静态成员和访问修饰符\n\n### 本周练习\n1. 创建学生类并实现基本操作\n2. 设计银行账户类\n3. 实现图书管理系统基础类\n\n## 第 4 周：面向对象进阶\n### 学习重点\n- 继承的概念与实现\n- 方法重写与super关键字\n- 多态性与抽象类\n\n### 本周练习\n1. 创建动物类继承体系\n2. 实现图形面积计算系统\n3. 设计员工薪资管理系统\n\n## 第 5 周：异常处理与集合框架\n### 学习重点\n- 异常处理机制（try-catch-finally）\n- 自定义异常\n- ArrayList、LinkedList、HashMap等集合类\n\n### 本周练习\n1. 实现带异常处理的文件读写程序\n2. 创建学生成绩管理系统（使用集合）\n3. 设计购物车系统\n\n## 第 6 周：IO流与文件操作\n### 学习重点\n- 字节流与字符流的区别\n- FileInputStream/FileOutputStream\n- BufferedReader/BufferedWriter\n- 序列化与反序列化\n\n### 本周练习\n1. 实现文本文件复制程序\n2. 创建日志记录系统\n3. 设计配置文件读写工具\n\n## 第 7 周：多线程编程\n### 学习重点\n- 线程创建方式（继承Thread、实现Runnable）\n- 线程同步与锁机制\n- 线程通信（wait/notify）\n\n### 本周练习\n1. 实现生产者消费者模型\n2. 创建多线程下载器\n3. 设计线程安全的计数器\n\n## 第 8 周：项目实战与复习\n### 学习重点\n- 综合运用所学知识\n- 代码调试与优化技巧\n- 面试常见问题准备\n\n### 本周练习\n1. 开发简易图书管理系统（完整MVC架构）\n2. 实现在线考试系统核心功能\n3. 准备技术面试问答集\n\n# Java（初学者）8 周练习与验收清单\n\n## 第 1 周验收\n- 能解释 JDK、JRE、JVM 的区别。\n- 能独立编写 HelloWorld、变量演示、四则运算程序。\n- 能说清楚基本数据类型及类型转换的使用场景。\n\n## 第 2 周验收\n- 能使用 if/else、switch、for、while 完成分支与循环练习。\n- 能独立完成九九乘法表、成绩评级、猜数字等 3 个练习。\n- 能通过调试定位循环边界错误。\n\n## 第 3 周验收\n- 能创建符合封装原则的学生类、银行账户类。\n- 能正确使用 this 和 static 关键字。\n- 能解释 private、public、protected 的访问控制效果。\n\n## 第 4 周验收\n- 能设计合理的类继承体系（如动物家族）。\n- 能正确重写父类方法并调用 super。\n- 能使用抽象类定义通用行为规范。\n\n## 第 5 周验收\n- 能为业务逻辑添加适当的异常处理。\n- 能选择合适的集合类存储不同类型的数据。\n- 能使用迭代器安全遍历集合。\n\n## 第 6 周验收\n- 能区分字节流与字符流的应用场景。\n- 能实现高效文件读写操作。\n- 能完成对象的序列化与反序列化。\n\n## 第 7 周验收\n- 能创建并启动多个线程执行并发任务。\n- 能使用 synchronized 解决线程安全问题。\n- 能实现线程间的协作与通信。\n\n## 第 8 周验收\n- 能独立完成一个中等复杂度的Java项目。\n- 能对代码进行性能分析和优化。\n- 能回答常见的Java面试问题。'''

# 简单的Markdown解析（仅处理标题和列表）
lines = md_content.split('\n')
for line in lines:
    line = line.strip()
    if not line:
        continue
    
    # 处理标题
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    # 处理列表项
    elif line.startswith('- '):
        p = doc.add_paragraph('', style='List Bullet')
        p.add_run(line[2:])
    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. ') or line.startswith('6. ') or line.startswith('7. ') or line.startswith('8. ') or line.startswith('9. '):
        p = doc.add_paragraph('', style='List Number')
        p.add_run(line[line.find('. ')+2:])
    else:
        # 普通段落
        doc.add_paragraph(line)

# 保存文档
output_path = 'java_learning_plan.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')