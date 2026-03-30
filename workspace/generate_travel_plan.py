from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_chinese_font(paragraph):
    """设置中文字体"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    rpr = run._r.get_or_add_rPr()
    east_asia = OxmlElement('w:eastAsia')
    east_asia.set(qn('w:val'), '微软雅黑')
    rpr.append(east_asia)

# 创建文档
doc = Document()

# 设置标题样式
style = doc.styles['Title']
font = style.font
font.size = Pt(24)
font.bold = True

# 添加标题
title = doc.add_heading('北京7天旅行计划', 0)
set_chinese_font(title)

doc.add_paragraph('\n')

# 添加行程概览
overview = doc.add_heading('行程概览', 1)
set_chinese_font(overview)

doc.add_paragraph('• 出发时间：下周1（具体日期需确认）')
doc.add_paragraph('• 出发地：上海')
doc.add_paragraph('• 目的地：北京')
doc.add_paragraph('• 旅行时长：7天')
doc.add_paragraph('• 特殊安排：1天工作日，1天购物日')

doc.add_paragraph('\n')

# 添加每日详细行程
schedule = doc.add_heading('每日详细行程', 1)
set_chinese_font(schedule)

# 行程内容（从getPlan获取）
plan_content = """1. 第1天：从上海出发抵达北京，入住酒店，适应环境，轻松休整
2. 第2天：工作日（安排会议/远程办公/商务活动）
3. 第3天：购物日（推荐三里屯、西单大悦城、王府井步行街）
4. 第4天：游览故宫（Forbidden City），建议上午入场，预留3-4小时，提前网上购票
5. 第5天：登八达岭长城，建议早出发，搭配明十三陵或奥林匹克公园
6. 第6天：颐和园 + 圆明园联游，感受皇家园林与历史遗迹
7. 第7天：自由活动/查漏补缺（可选南锣鼓巷、什刹海、国家博物馆或返程准备）"""

for line in plan_content.split('\n'):
    p = doc.add_paragraph(line)
    set_chinese_font(p)
    # 设置字体大小
    for run in p.runs:
        run.font.size = Pt(12)

doc.add_paragraph('\n')

# 添加景点简介
attractions = doc.add_heading('推荐景点简介', 1)
set_chinese_font(attractions)

doc.add_paragraph('• 故宫（Forbidden City）：明清两代皇宫，世界文化遗产，建议预留3-4小时参观')
doc.add_paragraph('• 八达岭长城：最著名的长城段落，建议早出发避开人流')
doc.add_paragraph('• 颐和园：中国现存规模最大、保存最完整的皇家园林')
doc.add_paragraph('• 圆明园：清代大型皇家园林，具有重要历史价值')
doc.add_paragraph('• 三里屯：北京时尚购物地标，国际品牌云集')
doc.add_paragraph('• 王府井步行街：北京著名商业街，传统与现代结合')

doc.add_paragraph('\n')

# 添加实用提示
tips = doc.add_heading('实用旅行提示', 1)
set_chinese_font(tips)

doc.add_paragraph('• 交通：北京地铁网络发达，建议购买交通卡或使用手机支付')
doc.add_paragraph('• 住宿：推荐住在东城区或朝阳区，交通便利')
doc.add_paragraph('• 天气：请提前查看北京天气预报，准备合适衣物')
doc.add_paragraph('• 购票：故宫、长城等热门景点需提前网上预约购票')
doc.add_paragraph('• 语言：大部分旅游区有英文标识，但学习简单中文问候语更有帮助')

doc.add_paragraph('\n')

# 添加结束语
end = doc.add_heading('祝您旅途愉快！', 2)
set_chinese_font(end)

# 保存文档
doc.save('北京7天旅行计划.docx')

print('旅行计划Word文档已生成：北京7天旅行计划.docx')