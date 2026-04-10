import json
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 读取 weather_report.md 文件（实际是JSON格式）
try:
    with open('weather_report.md', 'r', encoding='utf-8') as f:
        data = json.load(f)
except json.JSONDecodeError:
    # 如果不是JSON格式，尝试读取为纯文本
    with open('weather_report.md', 'r', encoding='utf-8') as f:
        content = f.read()
    data = {"raw_content": content}

# 创建新的Word文档
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)

# 添加标题
heading = doc.add_heading('厦门天气报告', level=1)
heading.runs[0].font.name = 'SimSun'
heading.runs[0].font.size = Pt(16)

# 添加天气信息
if 'city' in data:
    # 处理JSON格式的天气数据
    doc.add_paragraph(f'城市：{data.get("city", "未知")}')
    doc.add_paragraph(f'温度：{data.get("temperature", "未知")}°C')
    doc.add_paragraph(f'天气状况：{data.get("condition", "未知")}')
    doc.add_paragraph(f'湿度：{data.get("humidity", "未知")}%')
    doc.add_paragraph(f'风速：{data.get("windSpeed", "未知")}km/h')
    
    # 添加位置信息
    location = data.get('location', {})
    if location:
        doc.add_paragraph(f'经纬度：{location.get("lat", "未知")}, {location.get("lon", "未知")}')
else:
    # 处理纯文本内容
    doc.add_paragraph(data.get('raw_content', '无天气数据'))

# 保存文档
doc.save('weather_report.docx')

print('Word文档已成功生成：weather_report.docx')